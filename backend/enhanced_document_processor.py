"""
Enhanced Document Processing Engine for ProfileGPT
Supports multiple file formats with intelligent content extraction
"""

import os
import io
import re
import json
import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path

# Document processing libraries
try:
    import PyPDF2
    import pypdf
    from pypdf import PdfReader
except ImportError:
    PyPDF2 = None
    pypdf = None
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import markdown
    from bs4 import BeautifulSoup
    import html2text
except ImportError:
    markdown = None
    BeautifulSoup = None
    html2text = None

try:
    from unstructured.partition.auto import partition
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition.docx import partition_docx
    from unstructured.chunking.title import chunk_by_title
except ImportError:
    partition = None
    partition_pdf = None
    partition_docx = None
    chunk_by_title = None

try:
    import langdetect
    import textstat
except ImportError:
    langdetect = None
    textstat = None

logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    title: str
    content: str
    chunks: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    source_type: str
    file_type: str
    language: Optional[str] = None
    readability_score: Optional[float] = None

class EnhancedDocumentProcessor:
    """Enhanced document processor with support for multiple formats and intelligent extraction"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_text,
            'md': self._process_markdown,
            'html': self._process_html,
            'csv': self._process_csv,
        }

    def process_document(self, file_path: str, source_type: str = "document",
                        title: Optional[str] = None) -> ProcessedDocument:
        """Process a document file and extract structured content"""

        file_ext = Path(file_path).suffix.lower().lstrip('.')

        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Extract content using appropriate processor
        content, metadata = self.supported_formats[file_ext](file_path)

        # Generate title if not provided
        if not title:
            title = self._generate_title(content, Path(file_path).stem)

        # Detect language
        language = self._detect_language(content)

        # Calculate readability score
        readability_score = self._calculate_readability(content)

        # Create intelligent chunks
        chunks = self._create_intelligent_chunks(content, source_type, file_ext)

        # Enhanced metadata
        enhanced_metadata = {
            **metadata,
            'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            'word_count': len(content.split()),
            'char_count': len(content),
            'language': language,
            'readability_score': readability_score,
            'chunk_count': len(chunks),
        }

        return ProcessedDocument(
            title=title,
            content=content,
            chunks=chunks,
            metadata=enhanced_metadata,
            source_type=source_type,
            file_type=file_ext,
            language=language,
            readability_score=readability_score
        )

    def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Enhanced PDF processing with fallback methods"""
        content = ""
        metadata = {}

        # Method 1: Try unstructured library (best for complex layouts)
        if partition_pdf:
            try:
                elements = partition_pdf(file_path)
                content = "\n".join([str(element) for element in elements])
                metadata['extraction_method'] = 'unstructured'
                logger.info(f"PDF processed with unstructured library: {file_path}")
                return content, metadata
            except Exception as e:
                logger.warning(f"Unstructured PDF processing failed: {e}")

        # Method 2: Try pypdf (newest)
        if pypdf and PdfReader:
            try:
                with open(file_path, 'rb') as file:
                    reader = PdfReader(file)
                    metadata['page_count'] = len(reader.pages)

                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

                    metadata['extraction_method'] = 'pypdf'
                    logger.info(f"PDF processed with pypdf: {file_path}")
                    return content, metadata
            except Exception as e:
                logger.warning(f"pypdf processing failed: {e}")

        # Method 3: Fallback to PyPDF2
        if PyPDF2:
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    metadata['page_count'] = len(reader.pages)

                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

                    metadata['extraction_method'] = 'PyPDF2'
                    logger.info(f"PDF processed with PyPDF2: {file_path}")
                    return content, metadata
            except Exception as e:
                logger.error(f"All PDF processing methods failed: {e}")

        raise ValueError("No PDF processing libraries available")

    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Enhanced DOCX processing"""
        content = ""
        metadata = {}

        # Method 1: Try unstructured library
        if partition_docx:
            try:
                elements = partition_docx(file_path)
                content = "\n".join([str(element) for element in elements])
                metadata['extraction_method'] = 'unstructured'
                return content, metadata
            except Exception as e:
                logger.warning(f"Unstructured DOCX processing failed: {e}")

        # Method 2: python-docx
        if DocxDocument:
            try:
                doc = DocxDocument(file_path)
                paragraphs = []

                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)

                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            paragraphs.append(row_text)

                content = "\n".join(paragraphs)
                metadata['extraction_method'] = 'python-docx'
                metadata['paragraph_count'] = len([p for p in doc.paragraphs if p.text.strip()])
                metadata['table_count'] = len(doc.tables)

                return content, metadata
            except Exception as e:
                logger.error(f"DOCX processing failed: {e}")

        raise ValueError("No DOCX processing libraries available")

    def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process plain text files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    metadata = {'encoding': encoding, 'extraction_method': 'plain_text'}
                    return content, metadata
            except UnicodeDecodeError:
                continue

        raise ValueError("Could not decode text file with any supported encoding")

    def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()

        metadata = {'extraction_method': 'markdown'}

        if markdown:
            try:
                # Convert to HTML first, then to plain text for better structure preservation
                html = markdown.markdown(md_content)
                if html2text:
                    h = html2text.HTML2Text()
                    h.ignore_links = False
                    content = h.handle(html)
                else:
                    content = md_content  # Fallback to raw markdown
                metadata['conversion'] = 'markdown_to_html_to_text'
            except Exception:
                content = md_content
                metadata['conversion'] = 'raw_markdown'
        else:
            content = md_content

        return content, metadata

    def _process_html(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process HTML files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()

        metadata = {'extraction_method': 'html'}

        if BeautifulSoup and html2text:
            try:
                # Clean HTML and extract text
                soup = BeautifulSoup(html_content, 'html.parser')

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Convert to text
                h = html2text.HTML2Text()
                h.ignore_links = False
                content = h.handle(str(soup))
                metadata['conversion'] = 'beautifulsoup_html2text'
            except Exception:
                # Fallback: just strip HTML tags
                content = re.sub('<[^<]+?>', '', html_content)
                metadata['conversion'] = 'regex_strip'
        else:
            # Basic HTML tag removal
            content = re.sub('<[^<]+?>', '', html_content)
            metadata['conversion'] = 'basic_strip'

        return content, metadata

    def _process_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process CSV files"""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)

            # Convert to readable text format
            content_lines = []
            content_lines.append(f"Data Summary: {len(df)} rows, {len(df.columns)} columns")
            content_lines.append(f"Columns: {', '.join(df.columns.tolist())}")
            content_lines.append("")

            # Add sample data
            content_lines.append("Sample Data:")
            content_lines.append(df.head(10).to_string(index=False))

            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content_lines.append("\nNumeric Summary:")
                content_lines.append(df[numeric_cols].describe().to_string())

            content = "\n".join(content_lines)
            metadata = {
                'extraction_method': 'pandas',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist()
            }

        except ImportError:
            # Fallback: basic CSV processing
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()

            content = "CSV Data:\n" + "".join(lines[:100])  # First 100 lines
            metadata = {
                'extraction_method': 'basic_csv',
                'total_lines': len(lines)
            }

        return content, metadata

    def _detect_language(self, content: str) -> Optional[str]:
        """Detect the language of the content"""
        if langdetect and content.strip():
            try:
                return langdetect.detect(content)
            except Exception:
                return None
        return None

    def _calculate_readability(self, content: str) -> Optional[float]:
        """Calculate readability score"""
        if textstat and content.strip():
            try:
                return textstat.flesch_reading_ease(content)
            except Exception:
                return None
        return None

    def _generate_title(self, content: str, filename: str) -> str:
        """Generate a meaningful title from content or filename"""
        # Try to extract title from content
        lines = content.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if len(line) > 10 and len(line) < 100:
                # Look for title-like patterns
                if any(keyword in line.lower() for keyword in ['resume', 'cv', 'curriculum', 'profile']):
                    return line
                if line.isupper() or line.istitle():
                    return line

        # Fallback to filename
        return filename.replace('_', ' ').replace('-', ' ').title()

    def _create_intelligent_chunks(self, content: str, source_type: str,
                                 file_type: str) -> List[Dict[str, Any]]:
        """Create intelligent chunks based on document structure"""
        chunks = []

        # Try unstructured chunking first
        if chunk_by_title and file_type in ['pdf', 'docx']:
            try:
                # This requires the full unstructured processing pipeline
                # For now, we'll use our enhanced rule-based chunking
                pass
            except Exception:
                pass

        # Enhanced rule-based chunking
        chunks = self._rule_based_chunking(content, source_type)

        return chunks

    def _rule_based_chunking(self, content: str, source_type: str) -> List[Dict[str, Any]]:
        """Enhanced rule-based document chunking"""
        chunks = []

        # Split by sections first (headers, major breaks)
        section_patterns = [
            r'\n\s*#{1,6}\s+.+',  # Markdown headers
            r'\n\s*[A-Z][A-Z\s]{2,20}:?\n',  # ALL CAPS sections
            r'\n\s*[A-Z][a-z\s]{5,50}:?\n',  # Title case sections
            r'\n\s*\d+\.\s+[A-Z]',  # Numbered sections
            r'\n\s*[A-Z][A-Z\s]+\n',  # Short ALL CAPS titles
        ]

        # Try to split by sections
        sections = []
        current_section = ""

        for line in content.split('\n'):
            is_section_break = False
            for pattern in section_patterns:
                if re.match(pattern, '\n' + line):
                    is_section_break = True
                    break

            if is_section_break and current_section.strip():
                sections.append(current_section.strip())
                current_section = line
            else:
                current_section += '\n' + line

        if current_section.strip():
            sections.append(current_section.strip())

        # If no clear sections found, use the whole content
        if len(sections) <= 1:
            sections = [content]

        # Create chunks from sections
        chunk_id = 0
        for section in sections:
            if len(section.strip()) < 50:  # Skip very short sections
                continue

            # If section is still too long, split it further
            if len(section) > self.chunk_size:
                subsections = self._split_long_text(section)
                for subsection in subsections:
                    if len(subsection.strip()) >= 50:
                        chunks.append({
                            'chunk_id': chunk_id,
                            'text': subsection.strip(),
                            'section': f"Section {chunk_id + 1}",
                            'word_count': len(subsection.split()),
                            'char_count': len(subsection),
                        })
                        chunk_id += 1
            else:
                chunks.append({
                    'chunk_id': chunk_id,
                    'text': section.strip(),
                    'section': f"Section {chunk_id + 1}",
                    'word_count': len(section.split()),
                    'char_count': len(section),
                })
                chunk_id += 1

        return chunks

    def _split_long_text(self, text: str) -> List[str]:
        """Split long text into smaller chunks while preserving meaning"""
        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        chunks = []
        current_chunk = ""

        for paragraph in paragraphs:
            # If adding this paragraph would exceed chunk size, finalize current chunk
            if len(current_chunk) + len(paragraph) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # Handle any remaining overly long chunks
        final_chunks = []
        for chunk in chunks:
            if len(chunk) > self.chunk_size * 1.5:  # 50% tolerance
                # Split by sentences as last resort
                sentences = re.split(r'[.!?]+', chunk)
                sub_chunk = ""
                for sentence in sentences:
                    if len(sub_chunk) + len(sentence) > self.chunk_size and sub_chunk:
                        final_chunks.append(sub_chunk.strip())
                        sub_chunk = sentence
                    else:
                        sub_chunk += sentence + ". "
                if sub_chunk.strip():
                    final_chunks.append(sub_chunk.strip())
            else:
                final_chunks.append(chunk)

        return final_chunks

# Usage example
if __name__ == "__main__":
    processor = EnhancedDocumentProcessor()

    # Example usage
    try:
        doc = processor.process_document("sample.pdf", "resume", "John Doe Resume")
        print(f"Processed: {doc.title}")
        print(f"Content length: {len(doc.content)} chars")
        print(f"Chunks created: {len(doc.chunks)}")
        print(f"Language: {doc.language}")
        print(f"Readability: {doc.readability_score}")
    except Exception as e:
        print(f"Processing failed: {e}")